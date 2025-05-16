from docx import Document
from docx.shared import Inches
import os

def fill_template_with_image(template_path, output_path, text_data, image_data):
    doc = Document(template_path)

    # Replace text placeholders in document paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            for key, value in text_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    # Replace placeholders and insert images in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Handle image placeholders
                for key, image_info in image_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        # Support either (path) or (path, width, height)
                        if isinstance(image_info, tuple):
                            image_path = image_info[0]
                            width = image_info[1] if len(image_info) > 1 else 4.0
                            height = image_info[2] if len(image_info) > 2 else None
                        else:
                            image_path = image_info
                            width = 4.0
                            height = None

                        insert_image_in_cell(cell, image_path, width=width, height=height)

                # Replace text placeholders in each cell
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in text_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

   

    doc.save(output_path)
    print(f"✅ Saved filled document to: {output_path}")


def insert_image_in_cell(cell, image_path, width=4.0, height=None):
    """Clear cell and insert image with optional size"""
    cell.text = ""  # Clear the {{photo}} placeholder
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(
        image_path,
        width=Inches(width),
        height=Inches(height) if height else None
    )

# ✅ Example usage
if __name__ == "__main__":
    text_data = {
        "name": "Alice",
        "age": "30",
        "tenant_name": "John Doe",
        "tenant_phone": "+1 555-123-4567",
        "description": "Water leak from the ceiling in the living room.",
        "specific_area": "Apartment",
        "date": "2025-05-15",
        "license": "NYC123456",
        "incident_time": "14:30",
        "report_time": "15:35",
        "repored_to": "Supervisor",
        "reported_by": "Security"
    }

    image_data = {
        "photo": ("photo.png", 3.0, 3.0),         # width x height in inches
        "logo": ("logo.png", 3.0, 1.5)
    }

    fill_template_with_image(
        "template.docx",         # Template with {{name}}, {{photo}}, {{logo}}, etc.
        "filled_output.docx",    # Output file (will be overwritten)
        text_data,
        image_data
    )
