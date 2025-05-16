from flask import Flask, request, render_template, send_file
from docx import Document
from docx.shared import Inches
import os
from werkzeug.utils import secure_filename
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# Your Gmail credentials (App Password recommended)
EMAIL_ADDRESS = "hanadalimohamed1@gmail.com"
EMAIL_APP_PASSWORD = "ggyl dezk tdpp mqgh"

app = Flask(__name__)

UPLOAD_FOLDER = "uploaded"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg"}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def fill_template_with_image(template_path, output_path, text_data, image_data):
    doc = Document(template_path)
    print("📄 Starting document processing...")

    for para in doc.paragraphs:
        for run in para.runs:
            for key, value in text_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run.text:
                    print(f"🔤 Replacing placeholder: {placeholder} with: {value}")
                    run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, image_info in image_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        image_path, width, height = image_info
                        print(f"🖼️ Inserting image for {key} from {image_path}")
                        insert_image_in_cell(cell, image_path, width, height)
                        clear_placeholder_from_cell(cell, placeholder)

                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in text_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

    if os.path.exists(output_path):
        os.remove(output_path)
    doc.save(output_path)
    print(f"✅ Saved filled document to: {output_path}")

def insert_image_in_cell(cell, image_path, width=3.0, height=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width), height=Inches(height) if height else None)

def clear_placeholder_from_cell(cell, placeholder):
    for para in cell.paragraphs:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, "")

def send_email_with_attachment(subject, body, to, attachment_path):
    msg = MIMEMultipart()
    msg["From"] = EMAIL_ADDRESS
    msg["To"] = to
    msg["Subject"] = subject

    msg.attach(MIMEText(body, "plain"))

    with open(attachment_path, "rb") as f:
        part = MIMEApplication(f.read(), Name=os.path.basename(attachment_path))
        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(attachment_path)}"'
        msg.attach(part)

    server = smtplib.SMTP("smtp.gmail.com", 587)
    server.starttls()
    server.login(EMAIL_ADDRESS, EMAIL_APP_PASSWORD)
    server.send_message(msg)
    server.quit()


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        text_data = {
            "name": request.form["name"],
            "tenant_name": request.form.get("tenant_name", ""),
            "tenant_phone": request.form.get("tenant_phone", ""),
            "description": request.form.get("description", ""),
            "specific_area": request.form.get("specific_area", ""),
            "date": request.form.get("date", ""),
            "license": request.form.get("license", ""),
            "incident_time": request.form.get("incident_time", ""),
            "reported_time": request.form.get("incident_time", ""),
            "repored_to": request.form.get("repored_to", ""),
            "reported_by": request.form.get("reported_by", "")
        }

        image_data = {}
        photo_path = None
        output_path = "filled_output.docx"

        # Handle uploaded photo
        photo = request.files.get("photo")
        if photo and allowed_file(photo.filename):
            filename = secure_filename(photo.filename)
            photo_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            photo.save(photo_path)
            print(f"📥 Uploaded photo saved to: {photo_path}")
            image_data["photo"] = (photo_path, 3.0, 3.0)
        else:
            print("⚠️ No photo uploaded or invalid format")

        # Generate filled DOCX
        fill_template_with_image("template.docx", output_path, text_data, image_data)

        # Send email with attachment
        try:
            send_email_with_attachment(
                subject="Incident Report",
                body="Attached is the generated incident report.",
                to="hanadalimohamed1@gmail.com",  # Replace with desired recipient
                attachment_path=output_path
            )
            print("✅ Email sent successfully.")
        except Exception as e:
            print(f"❌ Email sending failed: {e}")

        # Clean up files
        if os.path.exists(output_path):
            os.remove(output_path)
        if photo_path and os.path.exists(photo_path):
            os.remove(photo_path)

        return "✅ Report submitted and emailed successfully!"

    return render_template("form.html")

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0')

