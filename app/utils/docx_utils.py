from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import os


def replace_placeholders_in_paragraph(paragraph, text_data):
    full_text = "".join(run.text for run in paragraph.runs)
    full_text = full_text.replace("\n", "").replace("\r", "")

    replaced = False
    for key, value in text_data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
            replaced = True

    if replaced and paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def insert_image_in_cell(cell, image_path, width=3.0, height=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(
        image_path,
        width=Inches(width),
        height=Inches(height) if height else None,
    )


def clear_placeholder_from_cell(cell, placeholder):
    for para in cell.paragraphs:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, "")


def fill_template_with_image(template_path, output_path, text_data, image_data):
    doc = Document(template_path)

    def set_font_style(run):
        font = run.font
        font.name = "Calibri"
        font.size = Pt(11)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, text_data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, image_info in image_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        image_path, width, height = image_info
                        insert_image_in_cell(cell, image_path, width, height)
                        clear_placeholder_from_cell(cell, placeholder)

                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in text_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
                                set_font_style(run)

    if os.path.exists(output_path):
        os.remove(output_path)

    doc.save(output_path)
